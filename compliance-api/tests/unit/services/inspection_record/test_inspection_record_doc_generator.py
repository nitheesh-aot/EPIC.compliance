"""Tests for inspection_record_doc_generator module helper methods."""

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from compliance_api.services.inspection_record.html_to_docx import (
    _add_formatted_list_item_text, _add_formatted_text_to_container, _add_formatted_text_to_table_cell,
    _add_html_paragraphs_to_cell, _add_html_to_container, _add_paragraph)


class TestAddFormattedTextToContainer:
    """Tests for _add_formatted_text_to_container function."""

    def _create_document_with_paragraph(self):
        """Create a document with a paragraph for testing."""
        doc = Document()
        para = doc.add_paragraph()
        return doc, para

    def test_plain_text_is_added(self):
        """Test that plain text is correctly added to paragraph."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Plain text content</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        assert para.text == "Plain text content"

    def test_bold_text_is_formatted(self):
        """Test that bold text has bold formatting applied."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Text with <strong>bold</strong> word</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        # Check that bold run exists and is marked bold
        bold_runs = [run for run in para.runs if run.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_italic_text_is_formatted(self):
        """Test that italic text has italic formatting applied."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Text with <em>italic</em> word</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        # Check that italic run exists and is marked italic
        italic_runs = [run for run in para.runs if run.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"

    def test_space_preserved_before_bold(self):
        """Test that space before bold text is preserved."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Before <strong>bold</strong> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        full_text = para.text
        assert full_text == "Before bold after"

    def test_space_preserved_after_bold(self):
        """Test that space after bold text is preserved."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Before <strong>bold</strong> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        full_text = para.text
        assert full_text == "Before bold after"

    def test_space_preserved_before_italic(self):
        """Test that space before italic text is preserved."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Before <em>italic</em> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        full_text = para.text
        assert full_text == "Before italic after"

    def test_multiple_formatted_words_preserve_spaces(self):
        """Test that multiple formatted words preserve all spaces."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Text <strong>bold</strong> and <em>italic</em> words</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        full_text = para.text
        assert full_text == "Text bold and italic words"

    def test_consecutive_formatted_elements_preserve_spaces(self):
        """Test that consecutive formatted elements preserve spaces between them."""
        doc, para = self._create_document_with_paragraph()

        html = "<p><strong>Bold</strong> <em>italic</em></p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        full_text = para.text
        assert full_text == "Bold italic"

    def test_b_tag_treated_as_bold(self):
        """Test that <b> tag is treated the same as <strong>."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Text with <b>bold</b> word</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        bold_runs = [run for run in para.runs if run.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_i_tag_treated_as_italic(self):
        """Test that <i> tag is treated the same as <em>."""
        doc, para = self._create_document_with_paragraph()

        html = "<p>Text with <i>italic</i> word</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        italic_runs = [run for run in para.runs if run.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"

    def test_whitespace_only_text_node_preserved(self):
        """Test that whitespace-only text nodes between elements are preserved."""
        doc, para = self._create_document_with_paragraph()

        # This HTML has a whitespace-only text node between </strong> and <em>
        html = "<p><strong>Bold</strong> <em>italic</em></p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_container(doc, p_element, Pt(11), first_para=para)

        # The space between Bold and italic must be preserved
        assert para.text == "Bold italic"
        assert " " in para.text


class TestAddFormattedListItemText:
    """Tests for _add_formatted_list_item_text function."""

    def _create_document_with_paragraph(self):
        """Create a document with a paragraph for testing."""
        doc = Document()
        para = doc.add_paragraph()
        return doc, para

    def _line_break_count(self, para):
        """Count <w:br/> elements in a paragraph."""
        return len(para._element.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'))

    def test_indented_list_item_has_no_line_breaks(self):
        """Test that template whitespace around a list item's span adds no <w:br/> (COMP-896)."""
        _, para = self._create_document_with_paragraph()

        # Same shape as the INSPECTION_SCOPE template: newlines/indentation around the span
        html = """<ol>
            <li class="editor-listitem" style="margin-bottom: 4px;">
                <span>Condition 1 - Documentation.</span>
            </li>
        </ol>"""
        soup = BeautifulSoup(html, "html.parser")
        li_element = soup.find("li")

        _add_formatted_list_item_text(para, li_element, Pt(11))

        assert para.text == "Condition 1 - Documentation."
        assert self._line_break_count(para) == 0

    def test_space_preserved_between_inline_elements(self):
        """Test that spaces between inline elements in a list item are preserved."""
        _, para = self._create_document_with_paragraph()

        html = "<li>Before <strong>bold</strong> after</li>"
        soup = BeautifulSoup(html, "html.parser")
        li_element = soup.find("li")

        _add_formatted_list_item_text(para, li_element, Pt(11))

        assert para.text == "Before bold after"

    def test_br_still_creates_line_break(self):
        """Test that an explicit <br> in a list item still produces a line break."""
        _, para = self._create_document_with_paragraph()

        html = "<li><span>First line</span><br><span>Second line</span></li>"
        soup = BeautifulSoup(html, "html.parser")
        li_element = soup.find("li")

        _add_formatted_list_item_text(para, li_element, Pt(11))

        assert self._line_break_count(para) == 1
        assert "First line" in para.text
        assert "Second line" in para.text


class TestAddFormattedTextToTableCell:
    """Tests for _add_formatted_text_to_table_cell function."""

    def _create_document_with_table_cell_paragraph(self):
        """Create a document with a table cell paragraph for testing."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        para = cell.paragraphs[0]
        return doc, cell, para

    def test_plain_text_is_added(self):
        """Test that plain text is correctly added to table paragraph."""
        _, cell, para = self._create_document_with_table_cell_paragraph()

        html = "<p>Plain text content</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_table_cell(cell, p_element, first_para=para)

        assert para.text == "Plain text content"

    def test_space_preserved_before_bold(self):
        """Test that space before bold text is preserved in table cells."""
        _, cell, para = self._create_document_with_table_cell_paragraph()

        html = "<p>Before <strong>bold</strong> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_table_cell(cell, p_element, first_para=para)

        full_text = para.text
        assert full_text == "Before bold after"

    def test_space_preserved_after_italic(self):
        """Test that space after italic text is preserved in table cells."""
        _, cell, para = self._create_document_with_table_cell_paragraph()

        html = "<p>Before <em>italic</em> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_table_cell(cell, p_element, first_para=para)

        full_text = para.text
        assert full_text == "Before italic after"

    def test_span_preserves_spaces(self):
        """Test that span elements preserve surrounding spaces."""
        _, cell, para = self._create_document_with_table_cell_paragraph()

        html = "<p>Before <span>span</span> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_table_cell(cell, p_element, first_para=para)

        full_text = para.text
        assert full_text == "Before span after"

    def test_whitespace_only_text_node_preserved(self):
        """Test that whitespace-only text nodes between elements are preserved."""
        _, cell, para = self._create_document_with_table_cell_paragraph()

        html = "<p><strong>Bold</strong> <em>italic</em></p>"
        soup = BeautifulSoup(html, "html.parser")
        p_element = soup.find("p")

        _add_formatted_text_to_table_cell(cell, p_element, first_para=para)

        assert para.text == "Bold italic"


class TestAddHtmlToContainer:
    """Tests for _add_html_to_container function."""

    def _create_document(self):
        """Create a document for testing."""
        return Document()

    def test_paragraph_with_bold_preserves_formatting_and_spaces(self):
        """Test that paragraphs with bold text preserve formatting and spaces."""
        doc = self._create_document()
        html = "<p>Text with <strong>bold</strong> word</p>"

        _add_html_to_container(doc, html)

        # Check the paragraph was added with proper spacing
        assert len(doc.paragraphs) >= 1
        para = doc.paragraphs[0]
        assert para.text == "Text with bold word"

    def test_paragraph_with_italic_preserves_formatting_and_spaces(self):
        """Test that paragraphs with italic text preserve formatting and spaces."""
        doc = self._create_document()
        html = "<p>Text with <em>italic</em> word</p>"

        _add_html_to_container(doc, html)

        para = doc.paragraphs[0]
        assert para.text == "Text with italic word"

    def test_spaces_preserved_around_formatted_text(self):
        """Test that spaces are preserved around formatted text."""
        doc = self._create_document()
        html = "<p>Before <strong>bold</strong> after</p>"

        _add_html_to_container(doc, html)

        para = doc.paragraphs[0]
        assert para.text == "Before bold after"

    def test_multiple_paragraphs_all_formatted_correctly(self):
        """Test that multiple paragraphs are all formatted correctly."""
        doc = self._create_document()
        html = """<p>First <strong>bold</strong> paragraph</p>
        <p>Second <em>italic</em> paragraph</p>"""

        _add_html_to_container(doc, html)

        # Check both paragraphs have correct spacing
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "First bold paragraph" in texts
        assert "Second italic paragraph" in texts

    def test_bold_formatting_is_applied(self):
        """Test that bold formatting is actually applied to the text."""
        doc = self._create_document()
        html = "<p>Text with <strong>bold</strong> word</p>"

        _add_html_to_container(doc, html)

        para = doc.paragraphs[0]
        bold_runs = [run for run in para.runs if run.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_italic_formatting_is_applied(self):
        """Test that italic formatting is actually applied to the text."""
        doc = self._create_document()
        html = "<p>Text with <em>italic</em> word</p>"

        _add_html_to_container(doc, html)

        para = doc.paragraphs[0]
        italic_runs = [run for run in para.runs if run.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"


class TestAddParagraph:
    """Tests for _add_paragraph function."""

    def _create_document(self):
        """Create a document for testing."""
        return Document()

    def test_plain_text_paragraph(self):
        """Test that plain text paragraphs are added correctly."""
        doc = self._create_document()

        html = "<p>Plain text paragraph</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_tag = soup.find("p")

        _add_paragraph(doc, p_tag, Pt(11))

        assert doc.paragraphs[0].text == "Plain text paragraph"

    def test_paragraph_with_bold_preserves_text_and_spaces(self):
        """Test that paragraphs with bold text preserve all text and spaces."""
        doc = self._create_document()

        html = "<p>Before <strong>bold</strong> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_tag = soup.find("p")

        _add_paragraph(doc, p_tag, Pt(11))

        # The full text should be preserved with spaces
        assert doc.paragraphs[0].text == "Before bold after"

    def test_paragraph_with_italic_preserves_text_and_spaces(self):
        """Test that paragraphs with italic text preserve all text and spaces."""
        doc = self._create_document()

        html = "<p>Before <em>italic</em> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_tag = soup.find("p")

        _add_paragraph(doc, p_tag, Pt(11))

        # The full text should be preserved with spaces
        assert doc.paragraphs[0].text == "Before italic after"

    def test_paragraph_with_bold_has_bold_formatting(self):
        """Test that paragraphs with bold text have bold formatting applied."""
        doc = self._create_document()

        html = "<p>Before <strong>bold</strong> after</p>"
        soup = BeautifulSoup(html, "html.parser")
        p_tag = soup.find("p")

        _add_paragraph(doc, p_tag, Pt(11))

        para = doc.paragraphs[0]
        bold_runs = [run for run in para.runs if run.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"


class TestAddHtmlParagraphsToCell:
    """Tests for _add_html_paragraphs_to_cell function."""

    def _create_table_cell(self):
        """Create a document with a table cell for testing."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        return doc, cell

    def test_simple_paragraph_in_cell(self):
        """Test that simple paragraphs are added to cells correctly."""
        _, cell = self._create_table_cell()
        html = "<p>Cell content</p>"

        _add_html_paragraphs_to_cell(cell, html)

        assert cell.paragraphs[0].text == "Cell content"

    def test_formatted_paragraph_preserves_spaces(self):
        """Test that formatted paragraphs preserve spaces in cells."""
        _, cell = self._create_table_cell()
        html = "<p>Before <strong>bold</strong> after</p>"

        _add_html_paragraphs_to_cell(cell, html)

        assert cell.paragraphs[0].text == "Before bold after"

    def test_italic_in_cell_preserves_spaces(self):
        """Test that italic text preserves surrounding spaces in cells."""
        _, cell = self._create_table_cell()
        html = "<p>Before <em>italic</em> after</p>"

        _add_html_paragraphs_to_cell(cell, html)

        assert cell.paragraphs[0].text == "Before italic after"
