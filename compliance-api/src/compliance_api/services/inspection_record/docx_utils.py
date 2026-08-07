"""Low-level DOCX/OOXML utility functions.

Provides helper functions for manipulating Word document XML elements
including hyperlinks, cell formatting, page numbers, and document settings.
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper around a new w:r element)
    new_run = OxmlElement('w:r')
    # Set the run's text
    rpr = OxmlElement('w:rPr')

    # Add color (blue) and underline for hyperlink style
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')  # hyperlink blue
    rpr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rpr.append(u)

    new_run.append(rpr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def _set_cell_background(cell, fill):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tc_para = tc.get_or_add_tcPr()

    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tc_borders.append(edge_el)
    tc_para.append(tc_borders)


def _add_page_number(run):
    """Add a page field code to display the current page number."""
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'PAGE'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _remove_compatibility_mode(doc):
    """Remove Word 2010 compatibility mode."""
    settings = doc.settings.element
    compatibility = settings.find(qn('w:compat'))
    if compatibility is not None:
        settings.remove(compatibility)


def _remove_cell_margins(cell):
    """Remove all margins from a table cell."""
    table_cell = cell._tc
    table_cell_pr = table_cell.get_or_add_tcPr()

    tc_margin = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:w'), '0')
        elem.set(qn('w:type'), 'dxa')
        tc_margin.append(elem)

    table_cell_pr.append(tc_margin)


def _remove_picture_spacing(run):
    """Set the wrap spacing around every inline picture in a run to zero.

    python-docx leaves the dist* attributes off <wp:inline>. Word reads the missing values as 0,
    but LibreOffice substitutes a 0.13" default that shows up as a gap beside the image.
    """
    for inline in run._r.findall('.//' + qn('wp:inline')):
        for side in ('distT', 'distB', 'distL', 'distR'):
            inline.set(side, '0')


def _set_empty_paragraph_spacing(doc):
    """Set line spacing to 0.5 for all empty paragraph blocks."""
    # Iterate through all paragraphs in the doc body
    for para in doc.paragraphs:
        if not para.text.strip():
            para.paragraph_format.line_spacing = 0.5

    # Check paragraphs inside table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip() and not _has_images(para):
                        para.paragraph_format.line_spacing = 0.5


def _has_images(para):
    """Check if a paragraph contains any images."""
    for run in para.runs:
        if run._element.xpath('.//pic:pic'):
            return True
    return False
