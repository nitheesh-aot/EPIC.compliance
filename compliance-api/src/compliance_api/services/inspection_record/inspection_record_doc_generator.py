"""DOCX Generator for Inspection Reports."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from flask import current_app

from .docx_utils import (
    _add_hyperlink, _add_page_number, _remove_cell_margins, _remove_compatibility_mode, _set_cell_background,
    _set_empty_paragraph_spacing)
from .html_to_docx import _add_html_paragraphs_to_cell, _add_html_to_container, init_list_numbering
from .image_utils import ImageDownloadError, ImageProcessingError, ImageTooLargeError, download_and_optimize_image


def _add_image_to_cell(cell, image_url, caption_text, error_prefix="image"):
    """Download, optimize, and add an image to a cell with caption.

    Uses streaming download and image optimization for memory efficiency.
    """
    if not image_url:
        return False

    image_stream = None
    try:
        image_stream = download_and_optimize_image(image_url)

        # Insert the image
        para = cell.add_paragraph()  # Add empty paragraph above
        para = cell.add_paragraph()
        run = para.add_run()
        run.add_picture(image_stream, width=Inches(4))

        # Insert caption below image
        if caption_text:
            caption_para = cell.add_paragraph()
            caption_run = caption_para.add_run(caption_text)
            caption_run.font.size = Pt(9)
        cell.add_paragraph()
        return True

    except (ImageDownloadError, ImageTooLargeError, ImageProcessingError) as e:
        # If image fails, show placeholder text
        para = cell.add_paragraph()
        para.text = f"[Failed to load {error_prefix}] {caption_text}"
        current_app.logger.info(f"Failed to load {error_prefix}: {caption_text}, from URL: {image_url}, error: {e}")
        return False

    finally:
        if image_stream is not None:
            image_stream.close()


def _add_photo(photo, cell):
    photo_url = photo.get('photo_url')
    caption_text = f"Photo {photo.get('photo_number', '')}. {photo.get('photo_caption', '')}"
    _add_image_to_cell(cell, photo_url, caption_text, error_prefix="image")


def _add_figure(figure, cell):
    figure_url = figure.get('figure_url')
    caption_text = (
        f"Figure {figure.get('figure_number', '')}. "
        f"{figure.get('figure_caption', '')}"
    )
    _add_image_to_cell(cell, figure_url, caption_text, error_prefix="figure")


def _add_inline_image_to_cell(cell, image_url, caption_text):
    """Add an inline image with simpler formatting for requirement/document images.

    Uses streaming download and image optimization for memory efficiency.
    """
    if not image_url:
        return

    image_stream = None
    try:
        cell.add_paragraph()  # Add spacing before image
        image_stream = download_and_optimize_image(image_url)

        img_para = cell.add_paragraph()
        run = img_para.add_run()
        run.add_picture(image_stream, width=Inches(4))

        if caption_text:
            caption_para = cell.add_paragraph()
            caption_run = caption_para.add_run(caption_text)
            caption_run.font.size = Pt(9)

    except (ImageDownloadError, ImageTooLargeError, ImageProcessingError) as e:
        error_para = cell.add_paragraph()
        error_para.text = f"[Failed to load image: {caption_text or 'unknown'}]"
        current_app.logger.info(f"Failed to load image: {caption_text or 'unknown'}, from URL: {image_url}, error: {e}")

    finally:
        if image_stream is not None:
            image_stream.close()


def _add_requirement_details_table(doc, req):
    req_table = doc.add_table(rows=0, cols=2)
    req_table.autofit = False
    req_table.columns[0].width = Inches(1.813)
    req_table.columns[1].width = Inches(5.439)
    req_table.style = 'Table Grid'

    # Requirement header and details
    source_details = req.get('requirement_source_details', [])
    last_description_is_table = False
    if source_details:
        row = req_table.add_row()
        cell = row.cells[0].merge(row.cells[1])
        para = cell.paragraphs[0]

        # Set Requirement header spacing
        para.paragraph_format.space_before = Inches(0.04)
        para.paragraph_format.space_after = Inches(0.02)

        ends_with_table = False
        for idx, source in enumerate(source_details):
            # Add requirement header for first source
            if idx == 0:
                run = para.add_run(f"Requirement {req.get('sort_order', '')}: ")
                para.paragraph_format.space_after = Inches(0.02)
                run.bold = True

            # Add requirement title
            req_title = source.get('requirement_title', '')
            run = para.add_run(req_title)
            run.bold = True

            # Add appendix if present
            if source.get('appendix_no'):
                run = para.add_run(f" (Appendix {source.get('appendix_no')})")

            para = cell.add_paragraph()

            # Add title if present (comes right after requirement title)
            if source.get('title'):
                run = para.add_run(source.get('title'))
                para.paragraph_format.space_after = Inches(0.02)
                run.bold = True

            # Add description
            if source.get("requirement_source_description"):
                _add_html_to_container(
                    cell,
                    source["requirement_source_description"],
                    clear_first=False,
                )
                ends_with_table = source.get("requirement_source_description", '').strip().endswith('</table>')

            # Add req source images
            for img in source.get('requirement_source_images', []):
                ends_with_table = False
                image_url = img.get('image_url')
                caption = img.get('original_file_name', '')
                _add_inline_image_to_cell(cell, image_url, caption)

            # Add document details (come after source description and images)
            for doc_group in source.get('requirement_documents', []):
                para = cell.add_paragraph()  # New paragraph for document group
                para = cell.add_paragraph()
                run = para.add_run(doc_group.get('document_title', ''))
                para.paragraph_format.space_after = Inches(0.02)
                run.bold = True

                # Check if first document has appendix
                documents = doc_group.get('documents', [])
                if documents and documents[0].get('appendix_no'):
                    ends_with_table = False
                    run = para.add_run(f" (Appendix {documents[0].get('appendix_no')})")
                    para = cell.add_paragraph()

                for document in documents:
                    ends_with_table = False
                    # Add section info
                    if document.get('section_number') or document.get('section_title'):
                        section_para = cell.add_paragraph()
                        section_text = ''
                        if document.get('section_number'):
                            section_text = f"Section {document.get('section_number')} "
                        if document.get('section_title'):
                            section_text += document.get('section_title')
                        run = section_para.add_run(section_text)
                        run.bold = True

                    # Add description
                    if document.get('description'):
                        ends_with_table = document.get("description", '').strip().endswith('</table>')
                        _add_html_to_container(
                            cell,
                            document.get('description'),
                            clear_first=False,
                        )
                    # Add document images
                    for img in document.get('document_images', []):
                        ends_with_table = False
                        image_url = img.get('image_url')
                        caption = img.get('original_file_name', '')
                        _add_inline_image_to_cell(cell, image_url, caption)
                    para = cell.add_paragraph()

            # Add spacing between multiple sources
            if idx < len(source_details) - 1:
                para = cell.add_paragraph()  # New paragraph for next source

            # Check if last description ends with a table
            if idx == len(source_details) - 1 and ends_with_table:
                last_description_is_table = True

    if last_description_is_table:
        cell.add_paragraph()

    # Inspection Details/Findings
    row = req_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    para = cell.paragraphs[0]
    # Set Header spacing
    para.paragraph_format.space_before = Inches(0.04)
    para.paragraph_format.space_after = Inches(0.08)
    run = para.add_run("Inspection Details:")
    run.bold = True
    _add_html_to_container(
        cell,
        req.get('requirement_findings', ''),
        clear_first=False
    )

    # Photos
    for photo in req.get('requirement_photos', []):
        _add_photo(photo, cell)

    # Figures
    for figure in req.get('requirement_figures', []):
        _add_figure(figure, cell)

    # Compliance Finding
    row = req_table.add_row()
    row.cells[0].text = "Compliance Finding"
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    _set_cell_background(row.cells[0], 'D9D9D9')
    row.cells[1].text = req.get('compliance_finding', '')

    # Enforcement Action
    row = req_table.add_row()
    row.cells[0].text = "Enforcement Action"
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    _set_cell_background(row.cells[0], 'BFBFBF')
    row.cells[1].text = req.get('enforcement_action', '')


def generate_inspection_report_docx(preview_data):
    """Generate a DOCX file from inspection report preview data."""
    doc = Document()
    _remove_compatibility_mode(doc)

    # Initialize list numbering definitions
    init_list_numbering(doc)

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.5)

    # Set font and spacing
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Inches(0.04)
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # List styles spacing
    list_bullet_style = doc.styles['List Bullet']
    list_bullet_style.paragraph_format.space_after = Inches(0.04)
    list_bullet_style.paragraph_format.left_indent = Inches(0.5)

    # Disable contextual spacing at style level
    para_properties = list_bullet_style._element.pPr
    if para_properties is not None:
        contextual_spacing = para_properties.find(qn('w:contextualSpacing'))
        if contextual_spacing is not None:
            para_properties.remove(contextual_spacing)

    list_number_style = doc.styles['List Number']
    list_number_style.paragraph_format.space_after = Inches(0.04)
    list_number_style.paragraph_format.left_indent = Inches(0.5)

    para_properties = list_number_style._element.pPr
    if para_properties is not None:
        contextual_spacing = para_properties.find(qn('w:contextualSpacing'))
        if contextual_spacing is not None:
            para_properties.remove(contextual_spacing)

    rpr = style._element.get_or_add_rPr()

    # Set East Asia font
    rfonts = rpr.rFonts
    rfonts.set(qn('w:eastAsia'), 'Calibri')

    # Add header with logo and title
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=Inches(7.25))
    header_table.autofit = False
    # Column widths
    header_table.columns[0].width = Inches(4.25)
    header_table.columns[1].width = Inches(3.0)
    for row in header_table.rows:
        row.cells[0].width = Inches(4.25)
        row.cells[1].width = Inches(3.0)
    if len(header.paragraphs) > 0:
        # Remove default paragraph
        p = header.paragraphs[0]
        p._element.getparent().remove(p._element)
    logo_cell = header_table.rows[0].cells[0]
    _remove_cell_margins(logo_cell)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Inches(0.18)

    logo_run = logo_para.add_run()
    logo_path = Path(__file__).parent / "assets" / "EAO_Logo.png"
    logo_run.add_picture(str(logo_path), width=Inches(4))

    # Add footer with page numbers
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer_run = footer_para.add_run()
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(51, 102, 153)  # #336699

    _add_page_number(footer_run)

    title_cell = header_table.rows[0].cells[1]
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    title_para = title_cell.paragraphs[0]
    title_para.text = "INSPECTION RECORD"
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_para.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(51, 102, 153)  # #336699

    # Extract data
    project_details = preview_data.get('project_details', {})
    inspection_details = preview_data.get('inspection_details', {})
    officer_details = preview_data.get('officer_details', {})
    department_details = preview_data.get('department_details', {})
    approval_info = preview_data.get('approval_info', {})
    version_date_info = preview_data.get('version_date_info', {})

    # Helper function to set header cell
    def set_header_cell(cell, text):
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True

    # First page table - Project and Inspection Information
    info_table = doc.add_table(rows=14, cols=4)
    info_table.style = 'Table Grid'

    # Row 1: Project Name and Inspection Status
    set_header_cell(info_table.cell(0, 0), "Project Name")
    info_table.cell(0, 1).text = project_details.get('name', '')
    set_header_cell(info_table.cell(0, 2), "Inspection Status")
    info_table.cell(0, 3).text = preview_data.get('ir_status', '')

    # Row 2: EA Certificate and Inspection No
    set_header_cell(info_table.cell(1, 0), project_details.get('certificate_label', 'EA Certificate #'))
    info_table.cell(1, 1).text = project_details.get('eac_certificate', '')
    set_header_cell(info_table.cell(1, 2), "Inspection No.")
    info_table.cell(1, 3).text = inspection_details.get('ir_number', '')

    # Row 3: Project Status and Inspection Start
    set_header_cell(info_table.cell(2, 0), "Project Status")
    info_table.cell(2, 1).text = project_details.get('project_state', '') or ''
    set_header_cell(info_table.cell(2, 2), "Inspection Start")
    info_table.cell(2, 3).text = inspection_details.get('start_date', '')

    # Row 4: Inspection Type and Initiation
    set_header_cell(info_table.cell(3, 0), "Inspection Type")
    info_table.cell(3, 1).text = inspection_details.get('inspection_type', '')
    set_header_cell(info_table.cell(3, 2), "Initiation")
    info_table.cell(3, 3).text = inspection_details.get('initiation', '')

    # Row 5: UTM
    set_header_cell(info_table.cell(4, 0), "UTM")
    merged_cell = info_table.cell(4, 1).merge(info_table.cell(4, 3))
    merged_cell.text = inspection_details.get('utm', '')

    # Row 6: Project Description
    set_header_cell(info_table.cell(5, 0), "Project Description")
    merged_cell = info_table.cell(5, 1).merge(info_table.cell(5, 3))
    merged_cell.text = inspection_details.get('project_description', '')

    # Row 7: Location Description
    set_header_cell(info_table.cell(6, 0), "Location Description")
    merged_cell = info_table.cell(6, 1).merge(info_table.cell(6, 3))
    merged_cell.text = inspection_details.get('location_description', '')

    # Row 8: Inspection Summary
    set_header_cell(info_table.cell(7, 0), "Inspection Summary")
    merged_cell = info_table.cell(7, 1).merge(info_table.cell(7, 3))

    merged_cell.text = ""

    if preview_data.get("inspection_scope"):
        _add_html_to_container(
            merged_cell,
            preview_data["inspection_scope"],
            clear_first=True,
        )

    if preview_data.get("preliminary_review_details"):
        _add_html_to_container(
            merged_cell,
            preview_data["preliminary_review_details"],
            clear_first=False,
        )

    if preview_data.get("finding_statement"):
        _add_html_to_container(
            merged_cell,
            preview_data["finding_statement"],
            clear_first=False,
        )

    # Row 9: In Attendance
    set_header_cell(info_table.cell(8, 0), "In Attendance")
    merged_cell = info_table.cell(8, 1).merge(info_table.cell(8, 3))
    attendance = officer_details.get('in_attendance', [])
    if attendance:
        merged_cell.text = '\n'.join(attendance) if isinstance(attendance, list) else str(attendance)
    else:
        merged_cell.text = 'N/A'

    # Row 10: Certificate Holder/Proponent
    set_header_cell(info_table.cell(9, 0), project_details.get('proponent_label', 'Certificate Holder'))
    merged_cell = info_table.cell(9, 1).merge(info_table.cell(9, 3))
    merged_cell.text = project_details.get('proponent', '')

    # Row 11: Mailing Address
    set_header_cell(info_table.cell(10, 0), "Mailing Address")
    merged_cell = info_table.cell(10, 1).merge(info_table.cell(10, 3))
    merged_cell.text = preview_data.get('mailing_address', 'None') or 'None'

    # Row 12: Inspecting Officer(s)
    set_header_cell(info_table.cell(11, 0), "Inspecting Officer(s)")
    merged_cell = info_table.cell(11, 1).merge(info_table.cell(11, 3))
    inspecting_officers = officer_details.get("inspecting_officers", [])

    if inspecting_officers:
        p = merged_cell.paragraphs[0]
        for i, officer in enumerate(inspecting_officers):
            if i > 0:
                p = merged_cell.add_paragraph()  # paragraph between officers
                p.paragraph_format.space_after = Inches(0.08)

            p.add_run(officer.get("name", ""))
            p.add_run("\n")
            p.add_run(officer.get("position", ""))

    # Row 13: Record Prepared By
    set_header_cell(info_table.cell(12, 0), "Record Prepared By")
    merged_cell = info_table.cell(12, 1).merge(info_table.cell(12, 3))
    record_prepared = officer_details.get('record_prepared_by', {})
    if record_prepared:
        merged_cell.text = f"{record_prepared.get('name', '')}\n{record_prepared.get('position', '')}"

    # Row 14: Record Approved By
    set_header_cell(info_table.cell(13, 0), "Record Approved By")
    merged_cell = info_table.cell(13, 1).merge(info_table.cell(13, 3))
    if approval_info:
        merged_cell.text = f"{approval_info.get('approved_by', '')}\n{approval_info.get('approved_by_position', '')}"

    # Inspection Details Section
    heading = doc.add_heading('INSPECTION DETAILS', level=1)
    heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
    heading.runs[0].font.size = Pt(11)

    # Requirement Inspection Details tables
    for req in preview_data.get('requirement_details', []):
        _add_requirement_details_table(doc, req)
        # Add spacing between requirements
        doc.add_paragraph()

    # Actions Required and Enforcement Summary table
    summary_table = doc.add_table(rows=0, cols=2)
    summary_table.style = 'Table Grid'

    # Actions Required by Certificate Holder
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.text = f"Actions Required by {project_details.get('proponent_label', 'Certificate Holder')}"
    cell.paragraphs[0].runs[0].font.bold = True
    _set_cell_background(cell, 'BFBFBF')

    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.text = ""

    if preview_data.get("action_required_by_rp"):
        _add_html_to_container(
            cell,
            preview_data["action_required_by_rp"],
        )
    else:
        para = cell.paragraphs[0]
        para.add_run("None at this time.")

    # Enforcement Summary
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.text = "Enforcement Summary"
    cell.paragraphs[0].runs[0].font.bold = True
    _set_cell_background(cell, 'BFBFBF')

    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])

    if preview_data.get("enforcement_summary"):
        _add_html_paragraphs_to_cell(
            cell,
            preview_data["enforcement_summary"],
        )
    else:
        cell.paragraphs[0].text = "None at this time."

    # Regulatory Considerations
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.text = "Regulatory Considerations"
    cell.paragraphs[0].runs[0].font.bold = True
    _set_cell_background(cell, 'BFBFBF')

    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    regulatory_consideration = preview_data.get('regulatory_consideration')
    cell.text = ""

    if regulatory_consideration and regulatory_consideration.get("findings"):
        _add_html_to_container(
            cell,
            regulatory_consideration["findings"],
        )

        # Add photos for regulatory considerations
        for photo in regulatory_consideration.get('photos', []):
            _add_photo(photo, cell)

        # Add figures for regulatory considerations
        for figure in regulatory_consideration.get('figures', []):
            _add_figure(figure, cell)
    else:
        para = cell.paragraphs[0]
        para.add_run("None at this time.")

    # Inspection Record Version Dates
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.text = "Inspection Record Version Dates"
    cell.paragraphs[0].runs[0].font.bold = True
    _set_cell_background(cell, 'BFBFBF')

    # Date Preliminary
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    para = cell.paragraphs[0]
    run = para.add_run("Date Preliminary")
    run.bold = True
    para = cell.add_paragraph()

    if version_date_info and version_date_info.get('preliminary_dates'):
        for date in version_date_info.get('preliminary_dates'):
            para.add_run(f"{date}\n")
        para.runs[-1].text = para.runs[-1].text.rstrip()
    else:
        para.add_run('n/a')

    # Date Issued
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    para = cell.paragraphs[0]
    run = para.add_run("Date Issued")
    run.bold = True
    para = cell.add_paragraph()

    if version_date_info and version_date_info.get('final_date'):
        para.add_run(version_date_info.get('final_date'))

    # Appendices
    if preview_data.get('appendices'):
        row = summary_table.add_row()
        cell = row.cells[0].merge(row.cells[1])
        cell.text = "Appendices:"
        cell.paragraphs[0].runs[0].font.bold = True
        _set_cell_background(cell, 'BFBFBF')

        row = summary_table.add_row()
        cell = row.cells[0].merge(row.cells[1])
        appendix_text = '\n'.join([f"Appendix {app.get('appendix_no', '')}: {app.get('document_title', '')}"
                                   for app in preview_data.get('appendices', [])])
        cell.text = appendix_text

    # Department details
    row = summary_table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    cell.text = "Environmental Assessment Office - Compliance & Enforcement Branch"
    cell.paragraphs[0].runs[0].font.bold = True
    _set_cell_background(cell, 'BFBFBF')

    row = summary_table.add_row()
    dept = department_details

    # Mailing Address
    para = row.cells[0].paragraphs[0]
    run = para.add_run("Mailing Address:")
    run.bold = True
    para.add_run(f"\n{dept.get('address_line1', '')}\n{dept.get('address_line2', '')}")

    # Contact Info
    para = row.cells[1].paragraphs[0]
    run = para.add_run("Phone: ")
    run.bold = True
    para.add_run(f"{dept.get('phone', '')}\n")

    run = para.add_run("Email: ")
    run.bold = True
    _add_hyperlink(para, dept.get('email', ''), f"mailto:{dept.get('email', '')}")
    para.add_run("\n")

    run = para.add_run("Website: ")
    run.bold = True
    _add_hyperlink(para, dept.get('website', ''), f"https://{dept.get('website', '')}")

    _set_empty_paragraph_spacing(doc)

    return doc
