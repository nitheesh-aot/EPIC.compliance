"""HTML to DOCX conversion utilities.

Converts HTML content (paragraphs, lists, tables, inline formatting)
to python-docx document elements. Supports nested lists with proper
numbering/bullet styles matching the web editor.
"""
import re

from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .docx_utils import set_cell_border


def _add_html_to_container(container, html_text, *, font_size=None, clear_first=True):
    """Add HTML content to a python-docx container (Document, Cell, or Paragraph)."""
    if not html_text:
        return

    soup = BeautifulSoup(html_text, "html.parser")

    # Clear the first paragraph if requested
    if clear_first and hasattr(container, 'paragraphs') and container.paragraphs:
        container.paragraphs[0].text = ""
        first_para_used = False
    else:
        first_para_used = True

    # if div, get its children
    children = soup.children
    if len(list(soup.children)) == 1 and list(soup.children)[0].name == 'div':
        children = list(soup.children)[0].children

    for element in children:
        if element.name == "p":
            # Get indentation from style
            indent_inches = _get_margin_left_inches(element)

            # Use the first empty paragraph if available
            if not first_para_used and hasattr(container, 'paragraphs') and container.paragraphs:
                para = container.paragraphs[0]
                # Handle text with inline formatting (italic, bold)
                if element.find(['i', 'em', 'strong', 'b', 'br']):
                    _add_formatted_text_to_container(container, element, font_size, first_para=para)
                else:
                    text = element.get_text(strip=True)
                    # Replace 2+ spaces with single space
                    text = re.sub(r' {2,}', ' ', text)
                    if text:
                        run = para.add_run(text)
                        if font_size:
                            run.font.size = font_size
                # Apply indentation
                if indent_inches > 0:
                    para.paragraph_format.left_indent = Inches(indent_inches)
                first_para_used = True
            else:
                _add_paragraph(container, element, font_size)

        elif element.name in ("ol", "ul"):
            _add_list(container, element, font_size)
            first_para_used = True  # Mark as used after adding list

        elif element.name == "table":
            add_html_table_to_container(container, element)
            first_para_used = True  # Mark as used after adding table


def _add_paragraph(container, p_tag, font_size):
    """Add a paragraph to the container, handling inline formatting and indentation."""
    # Blank paragraph (<p><br/></p>)
    para = container.add_paragraph()

    # Get indentation from style
    indent_inches = _get_margin_left_inches(p_tag)
    if indent_inches > 0:
        para.paragraph_format.left_indent = Inches(indent_inches)

    # Handle text with inline formatting (italic, bold, or br)
    if p_tag.find(['i', 'em', 'strong', 'b', 'br']):
        _add_formatted_text_to_container(container, p_tag, font_size, first_para=para)
    else:
        text = p_tag.get_text(strip=True)
        # Replace 2+ spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        if not text:
            return
        run = para.add_run(text)
        if font_size:
            run.font.size = font_size


def _add_formatted_text_to_container(container, element, font_size, first_para=None):
    """Add text with inline formatting (bold, italic) to a container, creating new paragraphs for <br> tags."""
    # Use the provided first paragraph or create a new one
    para = first_para if first_para is not None else container.add_paragraph()

    # Convert to list to check for remaining siblings
    children = list(element.children)

    for idx, child in enumerate(children):
        if isinstance(child, str):
            # Replace 2+ spaces with single space
            text = re.sub(r' {2,}', ' ', child)
            if text:
                run = para.add_run(text)
                if font_size:
                    run.font.size = font_size
        elif child.name in ['strong', 'b']:
            text = re.sub(r' {2,}', ' ', child.get_text())
            run = para.add_run(text)
            run.bold = True
            if font_size:
                run.font.size = font_size
        elif child.name in ['em', 'i']:
            text = re.sub(r' {2,}', ' ', child.get_text())
            run = para.add_run(text)
            run.italic = True
            if font_size:
                run.font.size = font_size
        elif child.name == 'br':
            # Only create a new paragraph if there's more content after this <br>
            remaining = children[idx + 1:]
            has_more_content = any(
                (isinstance(c, str) and c.strip()) or
                (hasattr(c, 'get_text') and c.get_text(strip=True))
                for c in remaining
            )
            if has_more_content:
                para = container.add_paragraph()
        else:
            text = re.sub(r' {2,}', ' ', child.get_text())
            if text:
                run = para.add_run(text)
                if font_size:
                    run.font.size = font_size


def _get_margin_left_inches(element):
    """Extract margin-left from element's style attribute and convert to inches."""
    style = element.get('style', '')
    if 'margin-left:' in style:
        # Extract the pixel value (e.g., "margin-left: 80px")
        match = re.search(r'margin-left:\s*(\d+)px', style)
        if match:
            px = int(match.group(1))
            # Convert pixels to inches (96 DPI)
            return px / 96
    return 0


def _add_html_paragraphs_to_cell(cell, html, font_size=Pt(11)):
    """Add HTML content as paragraphs to a table cell, handling inline formatting and indentation."""
    soup = BeautifulSoup(html, "html.parser")

    # Clear the first existing paragraph
    if cell.paragraphs:
        cell.paragraphs[0].text = ""
        first_para_used = False
    else:
        first_para_used = True

    # Find only direct child paragraphs, not nested ones
    # First, get the root element(s)
    root_elements = [el for el in soup.children if el.name == 'p']

    # If we have a wrapper <p>, get its children instead
    if len(root_elements) == 1 and root_elements[0].find_all('p', recursive=False):
        paragraphs = root_elements[0].find_all('p', recursive=False)
    else:
        paragraphs = soup.find_all('p', recursive=False)

    for p in paragraphs:
        # Handle empty <p><br/></p>
        text = p.get_text(strip=True)
        # Data has duplicate (sometimes four spaces) in a row
        # Replace 2+ spaces with single space
        text = re.sub(r' {2,}', ' ', text)

        # Get indentation from style
        indent_inches = _get_margin_left_inches(p)

        # Use the first empty paragraph if available
        if not first_para_used:
            para = cell.paragraphs[0]
            first_para_used = True
        else:
            para = cell.add_paragraph()

        # Apply indentation
        if indent_inches > 0:
            para.paragraph_format.left_indent = Inches(indent_inches)

        # Handle text with inline formatting (italic, bold, or br)
        if p.find(['i', 'em', 'strong', 'b', 'br']):
            _add_formatted_text_to_container(cell, p, font_size, first_para=para)
        else:
            run = para.add_run(text if text else "")
            run.font.size = font_size


def add_html_table_to_container(container, table_element):
    """Convert an HTML table to a docx table."""
    rows = table_element.find_all('tr')
    if not rows:
        return

    max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)

    # Create docx table
    docx_table = container.add_table(rows=len(rows), cols=max_cols)
    docx_table.style = 'Table Grid'
    docx_table.autofit = False

    # Set table width
    if hasattr(container, 'width'):
        parent_width_emu = container.width if isinstance(container.width, int) else container.width.emu
        table_width_emu = parent_width_emu - Inches(0.16).emu
    else:
        table_width_emu = Inches(7.09).emu

    # Apply width at XML level
    tbl = docx_table._element
    table_props = tbl.tblPr
    if table_props is None:
        table_props = OxmlElement('w:tblPr')
        tbl.insert(0, table_props)

    table_width = OxmlElement('w:tblW')
    table_width.set(qn('w:w'), str(int(table_width_emu / 635)))
    table_width.set(qn('w:type'), 'dxa')
    table_props.append(table_width)

    # Populate table
    for row_idx, tr in enumerate(rows):
        docx_row = docx_table.rows[row_idx]

        # Try to extract row height from style attribute
        style = tr.get('style', '')
        if 'height:' in style:
            # Extract height value (e.g., "66px")
            import re as regex
            height_match = regex.search(r'height:\s*(\d+)px', style)
            if height_match:
                height_px = int(height_match.group(1))
                # Convert pixels to inches (roughly 96 DPI)
                height_inches = height_px / 96
                docx_row.height = Inches(height_inches)

        cells = tr.find_all(['th', 'td'])
        for col_idx, cell in enumerate(cells):
            docx_cell = docx_row.cells[col_idx]

            # Clear default paragraph
            docx_cell.paragraphs[0].text = ""

            # Process cell content - handle paragraphs, lists, and formatting
            first_element = True
            for element in cell.children:
                if element.name == 'p':
                    # Use first paragraph or create new one
                    if first_element and docx_cell.paragraphs:
                        para = docx_cell.paragraphs[0]
                        first_element = False
                    else:
                        para = docx_cell.add_paragraph()

                    # Add formatted text to paragraph
                    _add_formatted_text_to_table_cell(docx_cell, element, first_para=para)

                elif element.name in ('ul', 'ol'):
                    # If list is first element, remove the empty default paragraph
                    if first_element and docx_cell.paragraphs and not docx_cell.paragraphs[0].text.strip():
                        p = docx_cell.paragraphs[0]._element
                        p.getparent().remove(p)
                    # Add list to cell
                    _add_list_to_table_cell(docx_cell, element, bullet_num_id=None, number_num_id=None)
                    first_element = False

            # Make header cells bold
            if cell.name == 'th':
                for para in docx_cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            # Set borders
            set_cell_border(
                docx_cell,
                top={"sz": 4, "val": "single", "color": "000000"},
                bottom={"sz": 4, "val": "single", "color": "000000"},
                start={"sz": 4, "val": "single", "color": "000000"},
                end={"sz": 4, "val": "single", "color": "000000"}
            )


def _add_formatted_text_to_table_cell(cell, p_element, first_para=None):
    """Add text with formatting (bold, italic) to a table cell, creating new paragraphs for <br> tags."""
    # Use the provided first paragraph or create a new one
    para = first_para if first_para is not None else cell.add_paragraph()

    # Convert to list to check for remaining siblings
    children = list(p_element.children)

    for idx, child in enumerate(children):
        if isinstance(child, str):
            text = re.sub(r' {2,}', ' ', child)
            if text:
                run = para.add_run(text)
                run.font.size = Pt(11)
        elif child.name in ['strong', 'b']:
            text = re.sub(r' {2,}', ' ', child.get_text())
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
        elif child.name in ['em', 'i']:
            text = re.sub(r' {2,}', ' ', child.get_text())
            run = para.add_run(text)
            run.italic = True
            run.font.size = Pt(11)
        elif child.name == 'span':
            text = re.sub(r' {2,}', ' ', child.get_text())
            if text:
                run = para.add_run(text)
                run.font.size = Pt(11)
        elif child.name == 'br':
            # Only create a new paragraph if there's more content after this <br>
            remaining = children[idx + 1:]
            has_more_content = any(
                (isinstance(c, str) and c.strip()) or
                (hasattr(c, 'get_text') and c.get_text(strip=True))
                for c in remaining
            )
            if has_more_content:
                para = cell.add_paragraph()


def _get_multilevel_number_num_id(document):
    """Create and return a numId for a multi-level numbered list with different formats per level."""
    numbering_part = document.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    # Count existing abstract nums and nums
    existing_abstract_nums = numbering.findall(qn('w:abstractNum'))
    existing_nums = numbering.findall(qn('w:num'))

    # Create a new abstract numbering definition
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num_id = len(existing_abstract_nums) + 200  # Different offset from bullets
    abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))

    # Multi-level type
    multi_level = OxmlElement('w:multiLevelType')
    multi_level.set(qn('w:val'), 'hybridMultilevel')
    abstract_num.append(multi_level)

    # Define number formats for each level same as in web editor
    num_formats = [
        'decimal',      # Level 0: 1, 2, 3
        'upperLetter',  # Level 1: A, B, C
        'lowerLetter',  # Level 2: a, b, c
        'upperRoman',   # Level 3: I, II, III
        'lowerRoman',   # Level 4: i, ii, iii

    ]

    # Level text patterns (the format string, %1 = level 1 number, etc.)
    lvl_texts = [
        '%1.',   # Level 0: 1.
        '%2.',   # Level 1: A.
        '%3.',   # Level 2: a.
        '%4.',   # Level 3: I.
        '%5.',   # Level 4: i.
    ]

    for level in range(9):
        idx = level % len(lvl_texts)  # Cycle through defined formats
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # Number format
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), num_formats[idx])
        lvl.append(num_fmt)

        # Level text (e.g., "%1." for "1.")
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), lvl_texts[idx])
        lvl.append(lvl_text)

        # Justification
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        # Paragraph properties (indentation)
        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        left_indent = 720 + (level * 720)  # 720 twips = 0.5 inch
        hanging = 360  # 360 twips = 0.25 inch
        ind.set(qn('w:left'), str(left_indent))
        ind.set(qn('w:hanging'), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)

        abstract_num.append(lvl)

    # Add the abstract numbering to the numbering part
    numbering.append(abstract_num)

    # Create a num element that references this abstract numbering
    num = OxmlElement('w:num')
    new_num_id = len(existing_nums) + 1
    num.set(qn('w:numId'), str(new_num_id))

    abstract_num_id_ref = OxmlElement('w:abstractNumId')
    abstract_num_id_ref.set(qn('w:val'), str(abstract_num_id))
    num.append(abstract_num_id_ref)

    numbering.append(num)

    return new_num_id


def _get_multilevel_bullet_num_id(document):
    """Create and return a numId for a multi-level bullet list with different bullet styles."""
    numbering_part = document.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    # Count existing abstract nums and nums by iterating
    existing_abstract_nums = numbering.findall(qn('w:abstractNum'))
    existing_nums = numbering.findall(qn('w:num'))

    # Create a new abstract numbering definition
    abstract_num = OxmlElement('w:abstractNum')
    # Use a high number to avoid conflicts
    abstract_num_id = len(existing_abstract_nums) + 100
    abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))

    # Multi-level type
    multi_level = OxmlElement('w:multiLevelType')
    multi_level.set(qn('w:val'), 'hybridMultilevel')
    abstract_num.append(multi_level)

    # Define bullet characters for each level same as in web editor
    bullet_chars = [
        '\u2022',  # Level 0: • (solid bullet)
        '\u25E6',  # Level 1: ◦ (hollow bullet)
        '\u25AA',  # Level 2+: ▪ (solid square)
    ]

    bullet_fonts = [
        'Symbol',
        'Symbol',
        'Wingdings',
    ]

    for level in range(9):
        idx = min(level, 2)  # Cap at level 2 style
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # Bullet format
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'bullet')
        lvl.append(num_fmt)

        # Bullet character
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), bullet_chars[idx])
        lvl.append(lvl_text)

        # Justification
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        # Paragraph properties (indentation)
        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        # Indentation increases with level
        left_indent = 720 + (level * 720)  # 720 twips = 0.5 inch
        hanging = 360  # 360 twips = 0.25 inch
        ind.set(qn('w:left'), str(left_indent))
        ind.set(qn('w:hanging'), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)

        # Run properties (font for the bullet)
        r_pr = OxmlElement('w:rPr')
        r_fonts = OxmlElement('w:rFonts')
        r_fonts.set(qn('w:ascii'), bullet_fonts[idx])
        r_fonts.set(qn('w:hAnsi'), bullet_fonts[idx])
        r_fonts.set(qn('w:hint'), 'default')
        r_pr.append(r_fonts)
        lvl.append(r_pr)

        abstract_num.append(lvl)

    # Add the abstract numbering to the numbering part
    numbering.append(abstract_num)

    # Create a num element that references this abstract numbering
    num = OxmlElement('w:num')
    new_num_id = len(existing_nums) + 1
    num.set(qn('w:numId'), str(new_num_id))

    abstract_num_id_ref = OxmlElement('w:abstractNumId')
    abstract_num_id_ref.set(qn('w:val'), str(abstract_num_id))
    num.append(abstract_num_id_ref)

    numbering.append(num)

    return new_num_id


def _add_list(container, list_tag, font_size, level=0, bullet_num_id=None, number_num_id=None):
    """Add a list to the container, handling nested lists recursively."""
    style = "List Number" if list_tag.name == "ol" else "List Bullet"

    document = container.part.document

    # Create multi-level numbering definitions at level 0 (lazy init)
    if bullet_num_id is None:
        bullet_num_id = _get_multilevel_bullet_num_id(document)
    if number_num_id is None:
        number_num_id = _get_multilevel_number_num_id(document)

    # Determine which numId to use based on current list type
    is_ordered = list_tag.name == "ol"
    style = "List Number" if is_ordered else "List Bullet"
    current_num_id = number_num_id if is_ordered else bullet_num_id

    for li in list_tag.find_all("li", recursive=False):
        nested_list = li.find(['ul', 'ol'], recursive=False)

        has_direct_content = any(
            (isinstance(child, str) and child.strip()) or
            (hasattr(child, 'name') and child.name not in ['ul', 'ol'] and child.get_text(strip=True))
            for child in li.children
        )

        if has_direct_content:
            para = container.add_paragraph(style=style)
            base_indent = 0.5 + (level * 0.5)
            para.paragraph_format.left_indent = Inches(base_indent)
            para.paragraph_format.hanging_indent = Inches(0.25)

            para_props = para._element.pPr
            num_props = para_props.numPr
            if num_props is None:
                num_props = para_props._add_numPr()

            # Set ilvl for the nesting level
            ilvl = num_props.ilvl
            if ilvl is None:
                ilvl = OxmlElement("w:ilvl")
                num_props.insert(0, ilvl)
            ilvl.set(qn("w:val"), str(level))

            # Set the numId based on the current list type (bullet or number)
            num_id_el = OxmlElement("w:numId")
            num_id_el.set(qn("w:val"), str(current_num_id))
            num_props.append(num_id_el)

            _add_formatted_list_item_text(para, li, font_size)

        if nested_list:
            # Pass both numIds so nested lists can switch between bullet/numbered
            _add_list(container, nested_list, font_size, level=level + 1,
                      bullet_num_id=bullet_num_id, number_num_id=number_num_id)


def _add_formatted_list_item_text(para, li_element, font_size):
    """Add formatted text from a list item to a paragraph, skipping nested lists."""
    for child in li_element.children:
        # Skip nested lists - they're handled separately
        if hasattr(child, 'name') and child.name in ['ul', 'ol']:
            continue

        if isinstance(child, str):
            text = re.sub(r' {2,}', ' ', child)
            if text:
                run = para.add_run(text)
                if font_size:
                    run.font.size = font_size
        elif child.name in ['strong', 'b']:
            text = re.sub(r' {2,}', ' ', child.get_text())
            run = para.add_run(text)
            run.bold = True
            if font_size:
                run.font.size = font_size
        elif child.name in ['em', 'i']:
            text = re.sub(r' {2,}', ' ', child.get_text())
            run = para.add_run(text)
            run.italic = True
            if font_size:
                run.font.size = font_size
        elif child.name == 'span':
            text = re.sub(r' {2,}', ' ', child.get_text())
            if text:
                run = para.add_run(text)
                if font_size:
                    run.font.size = font_size
        elif child.name == 'br':
            # Line break within list item - just add newline
            para.add_run('\n')
        else:
            # For any other element, get its text
            text = re.sub(r' {2,}', ' ', child.get_text())
            if text:
                run = para.add_run(text)
                if font_size:
                    run.font.size = font_size


def _add_list_to_table_cell(cell, list_element, level=0, bullet_num_id=None, number_num_id=None):
    """Add a list (ordered or unordered) to a table cell, handling nested lists and type switching."""
    document = cell.part.document

    # Create multi-level numbering definitions at level 0 (lazy init)
    if bullet_num_id is None:
        bullet_num_id = _get_multilevel_bullet_num_id(document)
    if number_num_id is None:
        number_num_id = _get_multilevel_number_num_id(document)

    # Determine which numId to use based on current list type
    is_ordered = list_element.name == "ol"
    style = "List Number" if is_ordered else "List Bullet"
    current_num_id = number_num_id if is_ordered else bullet_num_id

    for li in list_element.find_all('li', recursive=False):
        nested_list = li.find(['ul', 'ol'], recursive=False)

        has_direct_content = any(
            (isinstance(child, str) and child.strip()) or
            (hasattr(child, 'name') and child.name not in ['ul', 'ol'] and child.get_text(strip=True))
            for child in li.children
        )

        if has_direct_content:
            para = cell.add_paragraph(style=style)
            base_indent = 0.5 + (level * 0.5)
            para.paragraph_format.left_indent = Inches(base_indent)
            para.paragraph_format.hanging_indent = Inches(0.25)

            para_props = para._element.pPr
            num_props = para_props.numPr
            if num_props is None:
                num_props = para_props._add_numPr()

            # Set ilvl for the nesting level
            ilvl = num_props.ilvl
            if ilvl is None:
                ilvl = OxmlElement("w:ilvl")
                num_props.insert(0, ilvl)
            ilvl.set(qn("w:val"), str(level))

            # Set the numId based on the current list type (bullet or number)
            num_id_el = OxmlElement("w:numId")
            num_id_el.set(qn("w:val"), str(current_num_id))
            num_props.append(num_id_el)

            _add_formatted_list_item_text(para, li, Pt(11))

        if nested_list:
            # Pass both numIds so nested lists can switch between bullet/numbered
            _add_list_to_table_cell(cell, nested_list, level=level + 1,
                                    bullet_num_id=bullet_num_id, number_num_id=number_num_id)
